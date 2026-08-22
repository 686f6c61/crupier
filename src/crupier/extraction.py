"""Bounded local document extraction for model-ready file context."""

from __future__ import annotations

import csv
import io
import json
import shutil
import subprocess
import zipfile
from dataclasses import asdict, dataclass, field
from datetime import date, datetime, time
from pathlib import Path
from typing import Any, Protocol, runtime_checkable

from .errors import CrupierModelUnsupportedError


@dataclass(slots=True)
class ExtractedTable:
    name: str
    columns: list[str]
    rows: list[dict[str, str]]
    total_rows: int | None = None
    truncated: bool = False
    warnings: list[str] = field(default_factory=list)

    def to_dict(self, *, include_rows: bool = True) -> dict[str, Any]:
        data = asdict(self)
        if not include_rows:
            data.pop("rows", None)
        return data

    def to_prompt_text(self) -> str:
        lines = [
            f"#### Table: {self.name}",
            "columns: " + json.dumps(self.columns, ensure_ascii=False),
            "rows:",
        ]
        lines.extend(json.dumps(row, ensure_ascii=False, sort_keys=False) for row in self.rows)
        if self.truncated:
            lines.append("[table truncated by configured extraction limits]")
        return "\n".join(lines)


@dataclass(slots=True)
class ExtractedDocument:
    text: str = ""
    tables: list[ExtractedTable] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    extractor: str = ""
    truncated: bool = False

    def to_prompt_text(self) -> str:
        sections = [self.text.strip()] if self.text.strip() else []
        sections.extend(table.to_prompt_text() for table in self.tables)
        return "\n\n".join(sections)

    def to_dict(self, *, include_content: bool = False) -> dict[str, Any]:
        data: dict[str, Any] = {
            "extractor": self.extractor,
            "truncated": self.truncated,
            "warnings": self.warnings,
            "text_chars": len(self.text),
            "tables": [table.to_dict(include_rows=include_content) for table in self.tables],
        }
        if include_content:
            data["text"] = self.text
        return data


@runtime_checkable
class OCRAdapter(Protocol):
    name: str

    def extract(
        self,
        path: Path,
        *,
        max_chars: int,
        timeout_seconds: float | None = None,
    ) -> ExtractedDocument: ...


@dataclass(slots=True)
class TesseractOCRAdapter:
    languages: str = "eng"
    executable: str = "tesseract"
    name: str = "tesseract"

    def extract(
        self,
        path: Path,
        *,
        max_chars: int,
        timeout_seconds: float | None = None,
    ) -> ExtractedDocument:
        executable = shutil.which(self.executable)
        if not executable:
            raise CrupierModelUnsupportedError(
                f"OCR adapter {self.name!r} requires the {self.executable!r} executable."
            )
        command = [executable]
        if self.languages:
            command.extend(["-l", self.languages])
        command.extend(["--", str(path.resolve()), "stdout"])
        try:
            completed = subprocess.run(
                command,
                capture_output=True,
                check=False,
                text=True,
                timeout=timeout_seconds,
            )
        except subprocess.TimeoutExpired as exc:
            raise CrupierModelUnsupportedError(
                f"OCR extraction exceeded timeout_seconds={timeout_seconds}."
            ) from exc
        if completed.returncode != 0:
            detail = (completed.stderr or completed.stdout or "unknown OCR error").strip()
            raise CrupierModelUnsupportedError(f"Tesseract OCR failed: {detail[:1000]}")
        text = completed.stdout
        truncated = len(text) > max_chars
        return ExtractedDocument(
            text=text[:max_chars],
            warnings=["ocr_output_truncated"] if truncated else [],
            extractor=f"tesseract:{self.languages or 'default'}",
            truncated=truncated,
        )


def extract_spreadsheet(
    path: Path,
    *,
    max_file_bytes: int,
    max_rows: int,
    max_columns: int,
    max_cell_chars: int,
    max_sheets: int = 25,
) -> ExtractedDocument:
    _require_bounded_file(path, max_file_bytes=max_file_bytes)
    suffix = path.suffix.lower()
    if suffix in {".csv", ".tsv"}:
        return _extract_delimited(
            path,
            delimiter="\t" if suffix == ".tsv" else None,
            max_rows=max_rows,
            max_columns=max_columns,
            max_cell_chars=max_cell_chars,
        )
    if suffix == ".xlsx":
        return _extract_xlsx(
            path,
            max_rows=max_rows,
            max_columns=max_columns,
            max_cell_chars=max_cell_chars,
            max_sheets=max_sheets,
        )
    raise CrupierModelUnsupportedError(
        f"Spreadsheet format {suffix or '<unknown>'!r} is not executable. "
        "Supported formats are .csv, .tsv, and .xlsx."
    )


def extract_docx(
    path: Path,
    *,
    max_file_bytes: int,
    max_rows: int,
    max_columns: int,
    max_cell_chars: int,
    max_chars: int,
    max_tables: int = 25,
) -> ExtractedDocument:
    _require_bounded_file(path, max_file_bytes=max_file_bytes)
    if path.suffix.lower() != ".docx":
        raise CrupierModelUnsupportedError(
            f"Document format {path.suffix.lower() or '<unknown>'!r} is not executable. "
            "Basic local document extraction currently supports .docx."
        )
    _validate_office_zip(path, max_uncompressed_bytes=max(max_file_bytes * 20, 20_000_000))
    try:
        from docx import Document  # type: ignore[import-not-found]
        from docx.table import Table  # type: ignore[import-not-found]
        from docx.text.paragraph import Paragraph  # type: ignore[import-not-found]
    except ImportError as exc:
        raise CrupierModelUnsupportedError(
            "DOCX extraction requires the optional dependency: "
            "pip install 'crupier[documents]'."
        ) from exc

    try:
        document = Document(str(path))
    except Exception as exc:
        raise CrupierModelUnsupportedError(f"Could not parse DOCX document: {exc}") from exc
    text_parts: list[str] = []
    tables: list[ExtractedTable] = []
    warnings: list[str] = []
    remaining_chars = max(0, max_chars)
    truncated = False
    table_index = 0
    iterator = getattr(document, "iter_inner_content", None)
    blocks = iterator() if callable(iterator) else [*document.paragraphs, *document.tables]
    for block in blocks:
        if isinstance(block, Paragraph):
            value = block.text.strip()
            if not value:
                continue
            clipped = value[:remaining_chars]
            if clipped:
                text_parts.append(clipped)
                remaining_chars -= len(clipped)
            if len(clipped) < len(value):
                truncated = True
                break
        elif isinstance(block, Table):
            if table_index >= max(1, max_tables):
                warnings.append("document_tables_truncated")
                truncated = True
                break
            table_index += 1
            table = _table_from_matrix(
                name=f"table_{table_index}",
                matrix=[[cell.text for cell in row.cells] for row in block.rows],
                max_rows=max_rows,
                max_columns=max_columns,
                max_cell_chars=max_cell_chars,
            )
            tables.append(table)
    if truncated:
        warnings.append("document_text_truncated")
    if any(_docx_has_unsupported_content(part) for part in document.part.package.parts):
        warnings.append("docx_embedded_or_extended_content_may_be_omitted")
    return ExtractedDocument(
        text="\n\n".join(text_parts),
        tables=tables,
        warnings=warnings,
        extractor="python-docx",
        truncated=truncated or any(table.truncated for table in tables),
    )


def _extract_delimited(
    path: Path,
    *,
    delimiter: str | None,
    max_rows: int,
    max_columns: int,
    max_cell_chars: int,
) -> ExtractedDocument:
    raw = path.read_bytes()
    text, encoding_warning = _decode_delimited(raw)
    sample = text[:16_384]
    dialect: Any
    if delimiter is None:
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
    else:
        dialect = csv.excel_tab if delimiter == "\t" else csv.excel
    rows = csv.reader(io.StringIO(text, newline=""), dialect)
    matrix: list[list[str]] = []
    row_limit = max(1, max_rows) + 2
    try:
        for row in rows:
            matrix.append(row)
            if len(matrix) >= row_limit:
                break
    except csv.Error as exc:
        raise CrupierModelUnsupportedError(
            f"Could not parse delimited spreadsheet: {exc}"
        ) from exc
    table = _table_from_matrix(
        name=path.stem or "table",
        matrix=matrix,
        max_rows=max_rows,
        max_columns=max_columns,
        max_cell_chars=max_cell_chars,
    )
    warnings = [encoding_warning] if encoding_warning else []
    warnings.extend(table.warnings)
    return ExtractedDocument(
        tables=[table],
        warnings=_unique(warnings),
        extractor=f"csv:{'tab' if dialect.delimiter == chr(9) else dialect.delimiter}",
        truncated=table.truncated,
    )


def _extract_xlsx(
    path: Path,
    *,
    max_rows: int,
    max_columns: int,
    max_cell_chars: int,
    max_sheets: int,
) -> ExtractedDocument:
    _validate_office_zip(path, max_uncompressed_bytes=max(path.stat().st_size * 20, 20_000_000))
    try:
        from openpyxl import load_workbook  # type: ignore[import-untyped]
    except ImportError as exc:
        raise CrupierModelUnsupportedError(
            "XLSX extraction requires the optional dependency: "
            "pip install 'crupier[spreadsheets]'."
        ) from exc
    try:
        workbook = load_workbook(
            filename=str(path),
            read_only=True,
            data_only=True,
            keep_links=False,
        )
    except (OSError, ValueError, zipfile.BadZipFile) as exc:
        raise CrupierModelUnsupportedError(f"Could not parse XLSX workbook: {exc}") from exc
    tables: list[ExtractedTable] = []
    warnings: list[str] = []
    try:
        remaining_rows = max(1, max_rows)
        worksheets = workbook.worksheets
        for worksheet in worksheets[: max(1, max_sheets)]:
            matrix: list[list[str]] = []
            for row in worksheet.iter_rows(
                min_row=1,
                max_row=remaining_rows + 2,
                max_col=max(1, max_columns) + 1,
                values_only=True,
            ):
                matrix.append([_cell_text(value) for value in row])
            table = _table_from_matrix(
                name=worksheet.title,
                matrix=matrix,
                max_rows=remaining_rows,
                max_columns=max_columns,
                max_cell_chars=max_cell_chars,
            )
            if worksheet.sheet_state != "visible":
                table.warnings.append(f"sheet_state:{worksheet.sheet_state}")
            tables.append(table)
            warnings.extend(table.warnings)
            remaining_rows -= len(table.rows)
            if remaining_rows <= 0:
                if len(tables) < len(worksheets):
                    warnings.append("workbook_total_rows_truncated")
                break
        if len(worksheets) > max(1, max_sheets):
            warnings.append("workbook_sheets_truncated")
    finally:
        workbook.close()
    return ExtractedDocument(
        tables=tables,
        warnings=_unique(warnings),
        extractor="openpyxl:data_only",
        truncated=any(table.truncated for table in tables),
    )


def _table_from_matrix(
    *,
    name: str,
    matrix: list[list[Any]],
    max_rows: int,
    max_columns: int,
    max_cell_chars: int,
) -> ExtractedTable:
    row_cap = max(1, int(max_rows))
    column_cap = max(1, int(max_columns))
    cell_cap = max(1, int(max_cell_chars))
    if not matrix:
        return ExtractedTable(name=name, columns=[], rows=[], total_rows=0)
    raw_width = max((len(row) for row in matrix), default=0)
    width = min(raw_width, column_cap)
    first = [_clip_cell(value, cell_cap) for value in matrix[0][:width]]
    has_header = _looks_like_header(first, matrix[1: min(len(matrix), 6)])
    columns = _unique_columns(first if has_header else [], width)
    data_rows = matrix[1:] if has_header else matrix
    rows: list[dict[str, str]] = []
    cell_truncated = False
    for source_row in data_rows[:row_cap]:
        values: list[str] = []
        for value in source_row[:width]:
            raw = _cell_text(value)
            clipped = raw[:cell_cap]
            cell_truncated = cell_truncated or len(clipped) < len(raw)
            values.append(clipped)
        values.extend([""] * (width - len(values)))
        rows.append(dict(zip(columns, values, strict=True)))
    truncated = len(data_rows) > row_cap or raw_width > column_cap or cell_truncated
    warnings: list[str] = []
    if len(data_rows) > row_cap:
        warnings.append("table_rows_truncated")
    if raw_width > column_cap:
        warnings.append("table_columns_truncated")
    if cell_truncated:
        warnings.append("table_cells_truncated")
    return ExtractedTable(
        name=name,
        columns=columns,
        rows=rows,
        total_rows=len(data_rows) if len(matrix) <= row_cap + 1 else None,
        truncated=truncated,
        warnings=warnings,
    )


def _decode_delimited(raw: bytes) -> tuple[str, str | None]:
    for encoding in ("utf-8-sig", "utf-8"):
        try:
            return raw.decode(encoding), None
        except UnicodeDecodeError:
            continue
    return raw.decode("cp1252", errors="replace"), "spreadsheet_decoded_as_cp1252"


def _looks_like_header(first: list[str], following: list[list[Any]]) -> bool:
    if not first or not any(value.strip() for value in first):
        return False
    if len({value.strip().casefold() for value in first if value.strip()}) != len(
        [value for value in first if value.strip()]
    ):
        return False
    first_non_numeric = sum(not _is_number(value) for value in first if value)
    later_values = [_cell_text(value) for row in following for value in row[: len(first)]]
    later_numeric = sum(_is_number(value) for value in later_values if value)
    return first_non_numeric == len([value for value in first if value]) and (
        later_numeric > 0 or all(value.replace("_", "").replace(" ", "").isalnum() for value in first)
    )


def _unique_columns(values: list[str], width: int) -> list[str]:
    columns: list[str] = []
    counts: dict[str, int] = {}
    for index in range(width):
        base = values[index].strip() if index < len(values) else ""
        base = base or f"column_{index + 1}"
        counts[base] = counts.get(base, 0) + 1
        columns.append(base if counts[base] == 1 else f"{base}_{counts[base]}")
    return columns


def _clip_cell(value: Any, max_chars: int) -> str:
    return _cell_text(value)[:max_chars]


def _cell_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime | date | time):
        return value.isoformat()
    return str(value)


def _is_number(value: str) -> bool:
    try:
        float(value.replace(",", "."))
    except (TypeError, ValueError):
        return False
    return True


def _require_bounded_file(path: Path, *, max_file_bytes: int) -> None:
    if not path.exists() or not path.is_file():
        raise CrupierModelUnsupportedError(f"File {str(path)!r} does not exist.")
    size = path.stat().st_size
    if size > max_file_bytes:
        raise CrupierModelUnsupportedError(
            f"File {str(path)!r} is {size} bytes, above max {max_file_bytes} bytes."
        )


def _validate_office_zip(path: Path, *, max_uncompressed_bytes: int) -> None:
    try:
        with zipfile.ZipFile(path) as archive:
            infos = archive.infolist()
            total = sum(info.file_size for info in infos)
            if len(infos) > 10_000 or total > max_uncompressed_bytes:
                raise CrupierModelUnsupportedError(
                    f"Office archive expands to {total} bytes across {len(infos)} entries, "
                    "above configured extraction safety limits."
                )
    except zipfile.BadZipFile as exc:
        raise CrupierModelUnsupportedError(f"Office document is not a valid ZIP container: {exc}") from exc


def _docx_has_unsupported_content(part: Any) -> bool:
    content_type = str(getattr(part, "content_type", "")).lower()
    return any(
        marker in content_type
        for marker in ("image", "comments", "footnotes", "endnotes", "oleobject")
    )


def _unique(values: list[str]) -> list[str]:
    return list(dict.fromkeys(value for value in values if value))
