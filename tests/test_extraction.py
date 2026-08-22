import subprocess
import zipfile
from datetime import date
from pathlib import Path
from types import SimpleNamespace

import pytest

from crupier.errors import CrupierModelUnsupportedError
from crupier.extraction import (
    ExtractedDocument,
    ExtractedTable,
    TesseractOCRAdapter,
    extract_docx,
    extract_spreadsheet,
)
from crupier.multimodal import normalize_file


def test_extracted_types_render_model_context_without_requiring_persistence():
    table = ExtractedTable(
        name="ledger",
        columns=["item", "amount"],
        rows=[{"item": "hosting", "amount": "42.50"}],
    )
    document = ExtractedDocument(
        text="Quarterly ledger",
        tables=[table],
        extractor="test",
    )

    assert '"amount": "42.50"' in document.to_prompt_text()
    assert "rows" not in document.to_dict()["tables"][0]
    assert document.to_dict(include_content=True)["tables"][0]["rows"][0]["item"] == "hosting"


def test_extracted_table_prompt_marks_truncation():
    table = ExtractedTable(name="bounded", columns=[], rows=[], truncated=True)

    assert table.to_prompt_text().endswith("[table truncated by configured extraction limits]")


def test_csv_classification_and_extraction_support_quotes_limits_and_encoding(tmp_path):
    csv_path = tmp_path / "orders.csv"
    csv_path.write_bytes(
        "name;notes;amount\nJosé;\"first line\nsecond line\";12,50\nAna;ok;9\nLuis;later;3\n".encode(
            "cp1252"
        )
    )

    assert normalize_file(csv_path).kind == "spreadsheet"
    document = extract_spreadsheet(
        csv_path,
        max_file_bytes=10_000,
        max_rows=2,
        max_columns=2,
        max_cell_chars=12,
    )

    table = document.tables[0]
    assert table.columns == ["name", "notes"]
    assert table.rows[0]["name"] == "José"
    assert "first line\ns" == table.rows[0]["notes"]
    assert table.truncated is True
    assert "spreadsheet_decoded_as_cp1252" in document.warnings
    assert "table_rows_truncated" in document.warnings
    assert "table_columns_truncated" in document.warnings
    assert "table_cells_truncated" in document.warnings


def test_tsv_without_header_uses_generated_column_names(tmp_path):
    path = tmp_path / "matrix.tsv"
    path.write_text("1\t2\n3\t4\n", encoding="utf-8")

    document = extract_spreadsheet(
        path,
        max_file_bytes=100,
        max_rows=10,
        max_columns=10,
        max_cell_chars=10,
    )

    assert document.tables[0].columns == ["column_1", "column_2"]
    assert document.tables[0].rows[0] == {"column_1": "1", "column_2": "2"}


def test_xlsx_extraction_is_data_only_bounded_and_reports_hidden_sheets(tmp_path):
    pytest.importorskip("openpyxl")
    from openpyxl import Workbook

    path = tmp_path / "budget.xlsx"
    workbook = Workbook()
    active = workbook.active
    active.title = "Budget"
    active.append(["item", "amount", "ignored"])
    active.append(["hosting", 42.5, "x"])
    active.append(["support", 15, "y"])
    hidden = workbook.create_sheet("Internal")
    hidden.sheet_state = "hidden"
    hidden.append(["secret"])
    hidden.append(["value"])
    workbook.save(path)

    document = extract_spreadsheet(
        path,
        max_file_bytes=1_000_000,
        max_rows=1,
        max_columns=2,
        max_cell_chars=20,
    )

    assert document.extractor == "openpyxl:data_only"
    assert document.tables[0].rows == [{"item": "hosting", "amount": "42.5"}]
    assert "table_rows_truncated" in document.tables[0].warnings
    assert "table_columns_truncated" in document.tables[0].warnings
    assert "workbook_total_rows_truncated" in document.warnings

    all_sheets = extract_spreadsheet(
        path,
        max_file_bytes=1_000_000,
        max_rows=10,
        max_columns=2,
        max_cell_chars=20,
    )
    assert "sheet_state:hidden" in all_sheets.tables[1].warnings

    one_sheet = extract_spreadsheet(
        path,
        max_file_bytes=1_000_000,
        max_rows=10,
        max_columns=2,
        max_cell_chars=20,
        max_sheets=1,
    )
    assert len(one_sheet.tables) == 1
    assert "workbook_sheets_truncated" in one_sheet.warnings


def test_docx_extraction_reads_paragraphs_and_tables_with_limits(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    path = tmp_path / "contract.docx"
    source = Document()
    source.add_paragraph("Renewal notice is thirty days.")
    table = source.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "clause"
    table.rows[0].cells[1].text = "owner"
    table.rows[1].cells[0].text = "renewal"
    table.rows[1].cells[1].text = "legal"
    table.rows[2].cells[0].text = "termination"
    table.rows[2].cells[1].text = "legal"
    source.save(path)

    document = extract_docx(
        path,
        max_file_bytes=1_000_000,
        max_rows=1,
        max_columns=2,
        max_cell_chars=20,
        max_chars=12,
    )

    assert document.text == "Renewal noti"
    assert document.truncated is True
    assert "document_text_truncated" in document.warnings
    assert document.tables == []


def test_docx_extraction_bounds_table_count(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    path = tmp_path / "tables.docx"
    source = Document()
    for index in range(3):
        table = source.add_table(rows=2, cols=1)
        table.cell(0, 0).text = "name"
        table.cell(1, 0).text = f"table-{index}"
    source.save(path)

    document = extract_docx(
        path,
        max_file_bytes=1_000_000,
        max_rows=10,
        max_columns=10,
        max_cell_chars=100,
        max_chars=1_000,
        max_tables=2,
    )

    assert len(document.tables) == 2
    assert "document_tables_truncated" in document.warnings
    assert document.truncated is True


def test_unsupported_and_invalid_office_files_fail_explicitly(tmp_path):
    unsupported = tmp_path / "legacy.xls"
    unsupported.write_bytes(b"legacy")
    corrupt = tmp_path / "broken.xlsx"
    corrupt.write_bytes(b"not-a-zip")

    with pytest.raises(CrupierModelUnsupportedError, match="Supported formats"):
        extract_spreadsheet(
            unsupported,
            max_file_bytes=100,
            max_rows=10,
            max_columns=10,
            max_cell_chars=10,
        )
    with pytest.raises(CrupierModelUnsupportedError, match="valid ZIP"):
        extract_spreadsheet(
            corrupt,
            max_file_bytes=100,
            max_rows=10,
            max_columns=10,
            max_cell_chars=10,
        )


def test_docx_rejects_wrong_suffix_missing_dependency_and_parse_failure(tmp_path, monkeypatch):
    from crupier import extraction

    wrong = tmp_path / "document.txt"
    wrong.write_bytes(b"plain")
    with pytest.raises(CrupierModelUnsupportedError, match="currently supports .docx"):
        extract_docx(
            wrong,
            max_file_bytes=100,
            max_rows=10,
            max_columns=10,
            max_cell_chars=10,
            max_chars=100,
        )

    document = tmp_path / "document.docx"
    with zipfile.ZipFile(document, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")

    real_import = __import__

    def missing_docx(name, *args, **kwargs):
        if name == "docx":
            raise ImportError("optional dependency unavailable")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr("builtins.__import__", missing_docx)
    with pytest.raises(CrupierModelUnsupportedError, match="optional dependency"):
        extract_docx(
            document,
            max_file_bytes=10_000,
            max_rows=10,
            max_columns=10,
            max_cell_chars=10,
            max_chars=100,
        )

    monkeypatch.setattr("builtins.__import__", real_import)
    monkeypatch.setattr(extraction, "_validate_office_zip", lambda *args, **kwargs: None)
    import docx

    monkeypatch.setattr(docx, "Document", lambda path: (_ for _ in ()).throw(ValueError("broken XML")))
    with pytest.raises(CrupierModelUnsupportedError, match="Could not parse DOCX.*broken XML"):
        extract_docx(
            document,
            max_file_bytes=10_000,
            max_rows=10,
            max_columns=10,
            max_cell_chars=10,
            max_chars=100,
        )


def test_docx_skips_blank_paragraphs(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    path = tmp_path / "paragraphs.docx"
    source = Document()
    source.add_paragraph("   ")
    source.add_paragraph("kept")
    source.save(path)

    document = extract_docx(
        path,
        max_file_bytes=1_000_000,
        max_rows=10,
        max_columns=10,
        max_cell_chars=10,
        max_chars=100,
    )

    assert document.text == "kept"


def test_delimited_parser_falls_back_from_sniffer_and_wraps_csv_errors(tmp_path, monkeypatch):
    from crupier import extraction

    path = tmp_path / "rows.csv"
    path.write_text("first,second\n1,2\n", encoding="utf-8")
    monkeypatch.setattr(
        extraction.csv.Sniffer,
        "sniff",
        lambda *args, **kwargs: (_ for _ in ()).throw(extraction.csv.Error("unknown dialect")),
    )
    assert extract_spreadsheet(
        path,
        max_file_bytes=100,
        max_rows=10,
        max_columns=10,
        max_cell_chars=10,
    ).extractor == "csv:,"

    class BrokenRows:
        def __iter__(self):
            return self

        def __next__(self):
            raise extraction.csv.Error("malformed row")

    monkeypatch.setattr(extraction.csv, "reader", lambda *args, **kwargs: BrokenRows())
    with pytest.raises(CrupierModelUnsupportedError, match="Could not parse.*malformed row"):
        extract_spreadsheet(
            path,
            max_file_bytes=100,
            max_rows=10,
            max_columns=10,
            max_cell_chars=10,
        )


def test_xlsx_reports_missing_dependency_and_parse_failure(tmp_path, monkeypatch):
    path = tmp_path / "sheet.xlsx"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
    real_import = __import__

    def missing_openpyxl(name, *args, **kwargs):
        if name == "openpyxl":
            raise ImportError("optional dependency unavailable")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr("builtins.__import__", missing_openpyxl)
    with pytest.raises(CrupierModelUnsupportedError, match="optional dependency"):
        extract_spreadsheet(
            path,
            max_file_bytes=10_000,
            max_rows=10,
            max_columns=10,
            max_cell_chars=10,
        )

    monkeypatch.setattr("builtins.__import__", real_import)
    import openpyxl

    monkeypatch.setattr(openpyxl, "load_workbook", lambda **kwargs: (_ for _ in ()).throw(ValueError("bad workbook")))
    with pytest.raises(CrupierModelUnsupportedError, match="Could not parse XLSX.*bad workbook"):
        extract_spreadsheet(
            path,
            max_file_bytes=10_000,
            max_rows=10,
            max_columns=10,
            max_cell_chars=10,
        )


def test_extraction_helpers_cover_empty_tables_headers_dates_and_file_limits(tmp_path):
    from crupier import extraction

    table = extraction._table_from_matrix(
        name="empty",
        matrix=[],
        max_rows=10,
        max_columns=10,
        max_cell_chars=10,
    )
    assert table.total_rows == 0
    assert extraction._looks_like_header([], []) is False
    assert extraction._looks_like_header(["same", "same"], [[1, 2]]) is False
    assert extraction._cell_text(date(2026, 8, 23)) == "2026-08-23"

    missing = tmp_path / "missing.csv"
    with pytest.raises(CrupierModelUnsupportedError, match="does not exist"):
        extract_spreadsheet(
            missing,
            max_file_bytes=10,
            max_rows=10,
            max_columns=10,
            max_cell_chars=10,
        )
    large = tmp_path / "large.csv"
    large.write_bytes(b"1234")
    with pytest.raises(CrupierModelUnsupportedError, match="above max 3 bytes"):
        extract_spreadsheet(
            large,
            max_file_bytes=3,
            max_rows=10,
            max_columns=10,
            max_cell_chars=10,
        )


def test_office_zip_rejects_excessive_entry_count(tmp_path, monkeypatch):
    from crupier import extraction

    class OversizedArchive:
        def __enter__(self):
            return self

        def __exit__(self, *args):
            return False

        def infolist(self):
            return [SimpleNamespace(file_size=0)] * 10_001

    monkeypatch.setattr(extraction.zipfile, "ZipFile", lambda path: OversizedArchive())
    with pytest.raises(CrupierModelUnsupportedError, match="10001 entries"):
        extraction._validate_office_zip(tmp_path / "large.docx", max_uncompressed_bytes=10)


def test_tesseract_adapter_success_timeout_missing_and_error(tmp_path, monkeypatch):
    from crupier import extraction

    image = tmp_path / "scan.png"
    image.write_bytes(b"image")
    adapter = TesseractOCRAdapter(languages="spa+eng")
    monkeypatch.setattr(extraction.shutil, "which", lambda executable: "/usr/bin/tesseract")
    monkeypatch.setattr(
        extraction.subprocess,
        "run",
        lambda command, **kwargs: SimpleNamespace(
            returncode=0,
            stdout="recognized text",
            stderr="",
        ),
    )

    result = adapter.extract(image, max_chars=10, timeout_seconds=5)

    assert result.text == "recognized"
    assert result.truncated is True
    assert result.extractor == "tesseract:spa+eng"

    monkeypatch.setattr(extraction.shutil, "which", lambda executable: None)
    with pytest.raises(CrupierModelUnsupportedError, match="requires"):
        adapter.extract(image, max_chars=10)

    monkeypatch.setattr(extraction.shutil, "which", lambda executable: "/usr/bin/tesseract")

    def timeout(*args, **kwargs):
        raise subprocess.TimeoutExpired(cmd="tesseract", timeout=1)

    monkeypatch.setattr(extraction.subprocess, "run", timeout)
    with pytest.raises(CrupierModelUnsupportedError, match="exceeded"):
        adapter.extract(image, max_chars=10, timeout_seconds=1)

    monkeypatch.setattr(
        extraction.subprocess,
        "run",
        lambda *args, **kwargs: SimpleNamespace(returncode=1, stdout="", stderr="bad image"),
    )
    with pytest.raises(CrupierModelUnsupportedError, match="bad image"):
        adapter.extract(image, max_chars=10)


def test_tesseract_command_never_uses_a_shell(tmp_path, monkeypatch):
    from crupier import extraction

    path = Path(tmp_path / "name with spaces.png")
    path.write_bytes(b"image")
    observed = {}
    monkeypatch.setattr(extraction.shutil, "which", lambda executable: "/opt/tesseract")

    def run(command, **kwargs):
        observed["command"] = command
        observed["kwargs"] = kwargs
        return SimpleNamespace(returncode=0, stdout="ok", stderr="")

    monkeypatch.setattr(extraction.subprocess, "run", run)

    TesseractOCRAdapter().extract(path, max_chars=100)

    assert observed["command"] == [
        "/opt/tesseract",
        "-l",
        "eng",
        "--",
        str(path.resolve()),
        "stdout",
    ]
    assert "shell" not in observed["kwargs"]


def test_ocr_command_uses_absolute_path_and_option_terminator(tmp_path, monkeypatch):
    from crupier import extraction

    monkeypatch.chdir(tmp_path)
    path = Path("-l.png")
    path.write_bytes(b"image")
    observed = {}
    monkeypatch.setattr(extraction.shutil, "which", lambda executable: "/opt/tesseract")

    def run(command, **kwargs):
        observed["command"] = command
        return SimpleNamespace(returncode=0, stdout="ok", stderr="")

    monkeypatch.setattr(extraction.subprocess, "run", run)

    TesseractOCRAdapter().extract(path, max_chars=100)

    separator = observed["command"].index("--")
    assert observed["command"][separator + 1] == str(path.resolve())
    assert Path(observed["command"][separator + 1]).is_absolute()
